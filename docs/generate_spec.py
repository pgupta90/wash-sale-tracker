from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Styles
def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def body(doc, text):
    return doc.add_paragraph(text)

def add_table_row(table, cells, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        if bg_color:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)
    return row

# Title
doc.add_heading('Wash Sale Prevention App — Design Specification', 0)
body(doc, 'Date: May 4, 2026')
body(doc, 'Stack: React (Vite) + Python (FastAPI) + SQLite')
doc.add_paragraph()

# ─── SECTION 1: Overview ───
heading(doc, '1. Overview', level=1)
body(doc, 'A single-user web app that helps prevent wash sale violations by syncing Robinhood trade history locally and allowing the user to look up all trades for a given symbol in the past 30 days before placing a new trade. The app does not auto-flag wash sales — it surfaces trade history so the user can make their own judgment.')
doc.add_paragraph()

# ─── SECTION 2: Architecture ───
heading(doc, '2. Architecture & Tech Stack', level=1)

heading(doc, 'Directory Structure', level=2)
doc.add_paragraph('/WashSaleApp\n  /backend          # FastAPI (Python)\n  /frontend         # React (Vite)\n  config.yaml       # Robinhood credentials + session token storage', style='Normal')

heading(doc, 'Backend', level=2)
items = [
    'Framework: FastAPI (Python)',
    'Robinhood integration: robin_stocks library',
    'Database: SQLite (local, stores synced trade history)',
    'Endpoints: POST /auth, POST /sync, GET /sync/status, GET /trades?symbol=&expiry=&strike=',
    'Auth flow: read username/password from config.yaml → Robinhood login → MFA prompt on CLI if needed → persist session token back to config.yaml',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading(doc, 'Frontend', level=2)
items = [
    'Framework: React + Vite',
    'Runs on localhost:5173, talks to backend at localhost:8000',
    'Three UI zones: sync status bar, search filters, results table',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading(doc, 'Data Flow', level=2)
steps = [
    '1. User runs python backend/main.py → server starts, auto-attempts login using config.yaml',
    '2. User runs npm run dev in /frontend → React app opens in browser',
    '3. User clicks "Sync Now" → backend fetches all stock + options trades from Robinhood → stores in SQLite',
    '4. User searches symbol (+ optional expiry/strike) → backend queries SQLite for last 30 days → returns results to React',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()

# ─── SECTION 3: Data Model ───
heading(doc, '3. Data Model', level=1)
body(doc, 'Single trades table in SQLite:')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Column'
hdr[1].text = 'Type'
hdr[2].text = 'Notes'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('id', 'TEXT', 'Robinhood order ID (primary key)'),
    ('symbol', 'TEXT', 'e.g. META'),
    ('platform', 'TEXT', 'Always "robinhood"'),
    ('trade_type', 'TEXT', '"option" or "stock"'),
    ('option_type', 'TEXT', '"call", "put", or null for stocks'),
    ('strategy', 'TEXT', '"single", "iron_condor", "call_spread", "put_spread"'),
    ('side', 'TEXT', '"buy" or "sell"'),
    ('expiration_date', 'DATE', 'null for stocks'),
    ('strike_price', 'DECIMAL', 'null for stocks'),
    ('trade_price', 'DECIMAL', 'Price per share/contract at execution'),
    ('quantity', 'DECIMAL', 'Number of shares or contracts'),
    ('status', 'TEXT', '"open" or "closed"'),
    ('executed_at', 'DATETIME', 'Trade execution timestamp'),
    ('synced_at', 'DATETIME', 'When this record was last pulled from Robinhood'),
]
for r in rows:
    add_table_row(table, list(r))

doc.add_paragraph()
body(doc, 'Search query logic: filter by symbol (required) + optional expiration_date + optional strike_price, where executed_at >= now − 30 days.')
doc.add_paragraph()

# ─── SECTION 4: UI Layout & Color Coding ───
heading(doc, '4. UI Layout & Color Coding', level=1)

heading(doc, 'Layout (single page, three zones)', level=2)
body(doc, 'Zone 1 — Sync Bar (top): Shows "Last synced: [datetime]" and a "Sync Now" button. During sync: button shows spinner and is disabled; timestamp replaced with "Syncing..."')
body(doc, 'Zone 2 — Search Filters: Symbol field (required), Expiry date field (optional), Strike price field (optional), Search button.')
body(doc, 'Zone 3 — Results Table: Shows count ("X trades in last 30 days"). Empty state: "No trades found for [SYMBOL] in the last 30 days".')

heading(doc, 'Results Table Columns (in order)', level=2)
body(doc, 'Symbol | Platform | Trade Type | Option Type | Strategy | Side | Expiry | Strike | Trade Price | Qty | Status | Date')

heading(doc, 'Color Coding', level=2)
doc.add_paragraph()

ctable = doc.add_table(rows=1, cols=2)
ctable.style = 'Table Grid'
chdr = ctable.rows[0].cells
chdr[0].text = 'Condition'
chdr[1].text = 'Visual Treatment'
for cell in chdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

color_rows = [
    ('Side = buy', 'Green text/badge'),
    ('Side = sell', 'Red text/badge'),
    ('Status = open', 'Bold row, white background'),
    ('Status = closed', 'Normal weight, light gray background'),
    ('Option Type = call', 'Blue badge'),
    ('Option Type = put', 'Orange badge'),
    ('Trade Type = stock (N/A)', 'Gray badge'),
    ('Strategy = iron_condor / spread', 'Purple badge on Strategy column'),
]
for r in color_rows:
    add_table_row(ctable, list(r))

doc.add_paragraph()

# ─── SECTION 5: Acceptance Criteria ───
heading(doc, '5. Acceptance Criteria', level=1)

heading(doc, 'Authentication', level=2)
acs = [
    'AC1: App reads config.yaml and logs in without manual input when credentials + valid session token are present.',
    'AC2: When session token is expired/missing, app prints MFA prompt to CLI, accepts input, and persists new token to config.yaml.',
    'AC3: Invalid credentials produce a clear error message; app does not crash.',
]
for ac in acs:
    doc.add_paragraph(ac, style='List Bullet')

heading(doc, 'Sync', level=2)
acs = [
    'AC4: "Last synced" timestamp is displayed in the UI and updates after each successful sync.',
    'AC5: Manual sync fetches all stock orders and all options orders (including legs of multi-leg strategies like iron condors).',
    'AC6: Re-syncing does not create duplicate records — existing rows are upserted by Robinhood order ID.',
    'AC7: Sync status (in-progress / complete / failed) is visible in the UI.',
]
for ac in acs:
    doc.add_paragraph(ac, style='List Bullet')

heading(doc, 'Search', level=2)
acs = [
    'AC8: Searching a symbol returns all trades for that symbol within the past 30 days.',
    'AC9: Adding expiry date filter narrows results to that expiry only.',
    'AC10: Adding strike price filter narrows results further to that exact strike.',
    'AC11: Results show all 12 columns: Symbol, Platform, Trade Type, Option Type, Strategy, Side, Expiry, Strike, Trade Price, Qty, Status, Date.',
    'AC12: Color coding distinguishes buy vs. sell and open vs. closed at a glance.',
    'AC13: Empty results show a clear "no trades found" message, not a blank table.',
]
for ac in acs:
    doc.add_paragraph(ac, style='List Bullet')

doc.add_paragraph()

# ─── SECTION 6: Test Plan ───
heading(doc, '6. Test Plan — Incremental Build Order', level=1)
body(doc, 'Each phase is independently testable before the next is built.')

heading(doc, 'Phase 1 — Auth (backend only, no UI)', level=2)
items = [
    'Run python backend/main.py with valid config → verify no MFA prompt, server starts.',
    'Clear session token from config → verify MFA prompt appears on CLI, new token written back.',
    'Use wrong password → verify error message, server still starts (unauthenticated state).',
    'Tool: curl http://localhost:8000/auth/status to check auth state.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading(doc, 'Phase 2 — Sync (backend only)', level=2)
items = [
    'Hit POST /sync via curl → verify SQLite trades table is populated.',
    'Inspect rows for at least one stock order and one options order.',
    'Re-run sync → verify row count stays the same (no duplicates).',
    'Verify iron condor legs appear as separate rows with strategy=iron_condor.',
    'Tool: sqlite3 backend/db.sqlite "SELECT trade_type, strategy, COUNT(*) FROM trades GROUP BY trade_type, strategy"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading(doc, 'Phase 3 — Search API (backend only)', level=2)
items = [
    'GET /trades?symbol=META → verify only META trades, within last 30 days.',
    'GET /trades?symbol=META&expiry=2025-06-20 → verify filtered results.',
    'GET /trades?symbol=META&expiry=2025-06-20&strike=500 → verify exact match.',
    'GET /trades?symbol=ZZZZ → verify empty array response, not an error.',
    'Tool: curl or Postman against localhost:8000.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading(doc, 'Phase 4 — React UI (frontend, with real backend)', level=2)
items = [
    'Sync status bar shows last synced time and "Sync Now" button.',
    'Clicking "Sync Now" triggers sync and updates timestamp.',
    'Search by symbol returns populated table with all 12 columns.',
    'Verify color coding: buy vs. sell visually distinct; open vs. closed visually distinct.',
    'Search with no results shows "no trades found" message.',
    'Tool: Manual browser testing + React DevTools.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading(doc, 'Phase 5 — End-to-End', level=2)
items = [
    'Full flow from cold start: no token → MFA prompt → sync → search → results.',
    'Verify a known options trade (e.g., a META put placed recently) appears correctly in results.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Save
out_path = '/Users/priya/Documents/claude-projects/WashSaleApp/docs/WashSaleApp-Design-Spec.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
